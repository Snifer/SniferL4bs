"""
Automatización para generar pruebas de manera aleatoria
Desarrollado por: Snifer
Website: www.sniferl4bs.com
Descripción:
Este script automatiza la generación de exámenes de manera aleatoria, asegurando que cada examen sea diferente. 
Para su funcionamiento, es necesario proporcionar un archivo CSV que contenga las preguntas en el siguiente formato:
1, Pregunta, Respuesta A, Respuesta B, Respuesta C

El script admite hasta cinco opciones de respuesta, siendo este el límite actual debido a la estructura del CSV. 
Funciona con base en un CSV preparado previamente.

Requerimientos:
- python-docx: Este módulo es necesario para manipular documentos de Word.
- template.docx: Este es el documento base que se utiliza y contiene el título del examen.
- preguntas.csv: Este es el archivo CSV que contiene las preguntas y las opciones de respuesta.

Ejemplo de Formato CSV:
1, ¿Cuál es la capital de Francia?, París, Londres, Berlín, Madrid, Roma

Nota:
Este script está limitado a la estructura especificada y no admite más de cinco opciones de respuesta por pregunta.

TODO: 

Formato de las pruebas generadas debe ser modificado previamente toca corregir para que se genere correctamente
"""

import csv
import random
import argparse
from docx import Document
from docx.shared import Pt

def read_questions(file_name):
    questions = []
    with open(file_name, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            questions.append(row)
    return questions

def select_random_questions(questions, num_questions):
    if num_questions > len(questions):
        raise ValueError("Error dude!")
    return random.sample(questions, num_questions)

def create_document(file_name, questions):
    doc = Document('template.docx')


    last_paragraph = doc.paragraphs[-1]
    
    for i, question in enumerate(questions, start=1):
        p = last_paragraph.insert_paragraph_before()
        p.add_run(f"{i}. {question[1]}\n").bold = True
        
        options = ['Ⓐ', 'Ⓑ', 'Ⓒ', 'Ⓓ', 'Ⓔ']
        for j, answer in enumerate(question[2:], start=0):
            p = doc.add_paragraph()
            p.add_run(f"{options[j]} ").bold = True
            p.add_run(answer)
    
        doc.add_paragraph()
    
    doc.save(file_name)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate random tests by SniferL4bs')
    parser.add_argument('questions_file', type=str, help='The CSV file containing the questions')
    parser.add_argument('num_exams', type=int, help='The number of exams to generate')
    parser.add_argument('num_questions', type=int, help='The number of questions per exam')

    args = parser.parse_args()

    questions_bank = read_questions(args.questions_file)
    
    if args.num_questions > len(questions_bank):
        print(f"Warning: The number of requested questions ({args.num_questions}) is greater than the available questions ({len(questions_bank)}). All available questions will be selected.")
        num_questions = len(questions_bank)
    else:
        num_questions = args.num_questions
    
    for doc_num in range(1, args.num_exams + 1):
        questions_selected = select_random_questions(questions_bank, num_questions)

        file_name = f'document_{doc_num}.docx'
        create_document(file_name, questions_selected)
        print(f'Test {file_name} created.')
